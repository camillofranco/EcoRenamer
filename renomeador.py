import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import customtkinter as ctk
from tkinter.ttk import Treeview
import openpyxl
from PIL import Image, ImageOps
import fitz  # PyMuPDF
import threading
import platform
import zipfile
import shutil
import subprocess
import webbrowser
import json
import urllib.request
import sys
import base64
import re
import tempfile
from concurrent.futures import ThreadPoolExecutor
import pdfplumber
from pdf2docx import Converter
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.pagesizes import landscape, A4

# Import opcional do Tesseract OCR (necessário para PDFs escaneados)
try:
    import pytesseract as _pytesseract_module
    _PYTESSERACT_OK = True
except ImportError:
    _PYTESSERACT_OK = False

VERSION = "1.4.5" # Versão estável sem IA
UPDATE_URL = "https://raw.githubusercontent.com/camillofranco/EcoRenamer/main/version.json"
REFS_URL = "https://github.com/camillofranco/EcoRenamer/releases"

def resource_path(relative_path):
    try:
        base_path = sys._MEIPASS
    except:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)

try:
    from pillow_heif import register_heif_opener
    register_heif_opener()
except ImportError:
    pass

# Configuração Global do Tema CTk
ctk.set_appearance_mode("System")  # Segue o sistema (Dark/Light)
ctk.set_default_color_theme("green") # Base verde do sistema

class ToolApp:
    def __init__(self, root):
        self.root = root
        self.root.title(f"EcoWave Pro v{VERSION} - Enterprise Edition")
        self.root.geometry("1000x800")
        self.root.minsize(900, 700) # Evita cortar telas pequenas
        
        # Paleta EcoWave
        self.c_primary = "#2E7D32" # Verde Forte
        self.c_secondary = "#4527A0" # Roxo Elegante
        
        # Variáveis Imagens
        self.img_folder = ctk.StringVar()
        self.excel_file = ctk.StringVar()
        self.digits = tk.IntVar(value=2)
        self.start_number = tk.IntVar(value=1)
        self.compress_var = ctk.BooleanVar(value=True)
        self.sort_order = ctk.StringVar(value="Decrescente (Z-A)")
        self.mapping = []
        
        # Variáveis PDF
        self.pdf_folder = ctk.StringVar()
        self.pdf_output_name = ctk.StringVar(value="Documento_Unificado.pdf")
        self.pdf_sort_order = ctk.StringVar(value="Crescente (A-Z)")
        self.processing = False
        self.pdf_files = []  # Lista de PDFs para merge
        
        self.create_widgets()

        
    def create_widgets(self):
        # 1. HEADER (Cabeçalho Branco Puro)
        # Em CTk, o Frame aceita cor sólida facilmente.
        self.header = ctk.CTkFrame(self.root, height=100, fg_color="#ffffff", corner_radius=0)
        self.header.pack(fill="x", side="top")
        self.header.pack_propagate(False)
        
        try:
            logo_img = Image.open(resource_path("logo_ecowave.png"))
            target_h = 75
            aspect = logo_img.width / logo_img.height
            logo_img = logo_img.resize((int(target_h * aspect), target_h), Image.Resampling.LANCZOS)
            self.logo_ctk = ctk.CTkImage(light_image=logo_img, dark_image=logo_img, size=(int(target_h * aspect), target_h))
            lbl_logo = ctk.CTkLabel(self.header, text="", image=self.logo_ctk)
            lbl_logo.pack(side="left", padx=40, pady=12)
        except Exception as e:
            ctk.CTkLabel(self.header, text="ECOWAVE PRO", font=ctk.CTkFont(size=28, weight="bold"), text_color=self.c_primary).pack(side="left", padx=40)
            
        badge = ctk.CTkFrame(self.header, fg_color="transparent")
        badge.pack(side="right", padx=40)
        ctk.CTkLabel(badge, text="ENTERPRISE EDITION", font=ctk.CTkFont(size=12, weight="bold"), text_color=self.c_secondary).pack()
        ctk.CTkLabel(badge, text=f"Build v{VERSION}", font=ctk.CTkFont(size=10), text_color="gray").pack()
        
        # 2. SELETOR DE ABAS PRINCIPAL (Muito mais bonito que o Notebook antigo)
        self.tabview = ctk.CTkTabview(self.root, corner_radius=10, segmented_button_selected_color=self.c_primary,
                                      segmented_button_selected_hover_color="#1B5E20")
        self.tabview.pack(fill="both", expand=True, padx=20, pady=20)
        
        self.create_img_tab(self.tabview.tab("📸  GESTÃO DE IMAGENS"))
        self.tabview.tab("📄  GESTÃO DE PDFS").columnconfigure(0, weight=1)
        self.tabview.tab("📄  GESTÃO DE PDFS").rowconfigure(1, weight=1)
        self.create_pdf_tab(self.tabview.tab("📄  GESTÃO DE PDFS"))
        self.create_utils_tab(self.tabview.tab("🛠️ UTILITÁRIOS"))
        
        # 3. BARRA DE STATUS INFERIOR
        self.status_bar = ctk.CTkFrame(self.root, height=40, corner_radius=0, fg_color=("gray90", "gray15"))
        self.status_bar.pack(fill="x", side="bottom")
        
        ctk.CTkLabel(self.status_bar, text=f"SO: {platform.system()} | Paralelismo Ativado (Motor Rápido)", 
                     font=ctk.CTkFont(size=11), text_color="gray").pack(side="left", padx=20)
        
        btn_upd = ctk.CTkButton(self.status_bar, text="♻️ Baixar Atualizações", command=self.check_for_updates,
                                fg_color=self.c_primary, text_color="white", font=ctk.CTkFont(size=12, weight="bold"), hover_color="#1B5E20", width=160, height=30)
        btn_upd.pack(side="right", padx=20, pady=5)
        
    def create_img_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(2, weight=1) # A tabela expande
        
        # --- BLOCO SUPERIOR (Configs) ---
        frame_config = ctk.CTkFrame(parent, fg_color="transparent")
        frame_config.grid(row=0, column=0, sticky="ew", pady=(0, 10))
        frame_config.columnconfigure(1, weight=1)
        
        # Linha 1: Imagens
        ctk.CTkLabel(frame_config, text="PASTA DE FOTOS:", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, sticky="w", pady=5)
        ctk.CTkEntry(frame_config, textvariable=self.img_folder, state="readonly", height=35).grid(row=0, column=1, sticky="ew", padx=10, pady=5)
        ctk.CTkButton(frame_config, text="Selecionar", command=self.select_img_folder, fg_color=self.c_secondary, hover_color="#311B92", width=100).grid(row=0, column=2, pady=5)
        
        # Linha 2: Excel
        ctk.CTkLabel(frame_config, text="EXCEL (Opcional):", font=ctk.CTkFont(weight="bold")).grid(row=1, column=0, sticky="w", pady=5)
        ctk.CTkEntry(frame_config, textvariable=self.excel_file, state="readonly", height=35).grid(row=1, column=1, sticky="ew", padx=10, pady=5)
        
        frame_b = ctk.CTkFrame(frame_config, fg_color="transparent")
        frame_b.grid(row=1, column=2, sticky="ew")
        ctk.CTkButton(frame_b, text="Buscar", command=self.select_excel, fg_color=self.c_secondary, width=70).pack(side="left", padx=(0,5))
        ctk.CTkButton(frame_b, text="Limpar", command=lambda: self.excel_file.set(""), fg_color="gray50", hover_color="gray30", width=70).pack(side="left")
        
        # Linha 3: Opções Mistas
        frame_opts = ctk.CTkFrame(frame_config, fg_color="transparent")
        frame_opts.grid(row=2, column=0, columnspan=3, sticky="w", pady=15)
        
        ctk.CTkLabel(frame_opts, text="Dígitos:", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=(0,5))
        # CTk não tem Spinbox nativo, usamos Entry (poderíamos criar um stepper)
        ctk.CTkEntry(frame_opts, textvariable=self.digits, width=50, justify="center").pack(side="left", padx=(0,20))
        
        ctk.CTkLabel(frame_opts, text="Nº Inicial:", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=(0,5))
        ctk.CTkEntry(frame_opts, textvariable=self.start_number, width=60, justify="center").pack(side="left", padx=(0,20))
        
        ctk.CTkLabel(frame_opts, text="Ordem Interna:", font=ctk.CTkFont(weight="bold")).pack(side="left", padx=(0,5))
        ctk.CTkComboBox(frame_opts, variable=self.sort_order, values=["Decrescente (Z-A)", "Crescente (A-Z)"], width=170).pack(side="left", padx=(0,20))
        
        ctk.CTkCheckBox(frame_opts, text="Comprimir em HD (Mais Rápido no App)", variable=self.compress_var, text_color=self.c_primary, font=ctk.CTkFont(weight="bold")).pack(side="left")

        # Botão Carregar (Grande)
        self.btn_load = ctk.CTkButton(frame_config, text="1. MONTAR ESTRUTURA DE NOMES", command=self.load_data, 
                                      fg_color=self.c_primary, text_color="white", font=ctk.CTkFont(size=14, weight="bold"), height=45)
        self.btn_load.grid(row=3, column=0, columnspan=3, sticky="ew", pady=(15, 0))


        # --- BLOCO CENTRAL (Tabela Tkinter Estilizada para CTk) ---
        frame_tree = ctk.CTkFrame(parent, fg_color="transparent")
        frame_tree.grid(row=2, column=0, sticky="nsew", pady=10)
        frame_tree.columnconfigure(0, weight=1)
        frame_tree.rowconfigure(0, weight=1)
        
        # É uma árvore Tkinter, mas daremos um ar moderno a ela alterando via ttk.Style
        style = ttk.Style()
        style.theme_use("clam")
        style.configure("Treeview", 
                        background="#ffffff" if ctk.get_appearance_mode() == "Light" else "#2b2b2b",
                        foreground="#000000" if ctk.get_appearance_mode() == "Light" else "#ffffff",
                        rowheight=35,
                        fieldbackground="#ffffff" if ctk.get_appearance_mode() == "Light" else "#2b2b2b",
                        font=("Helvetica", 11),
                        borderwidth=0)
        style.configure("Treeview.Heading", font=("Helvetica", 11, "bold"), background=self.c_primary, foreground="white", borderwidth=0)
        style.map("Treeview", background=[("selected", self.c_secondary)])
        
        cols = ("pos", "orig", "novo", "tam_orig", "tam_est")
        self.tree = Treeview(frame_tree, columns=cols, show="headings", selectmode="browse")
        for col in cols: self.tree.heading(col, text=col.upper())
        
        # Bindings Drag Drop
        self.tree.bind("<ButtonPress-1>", self.on_drag_start)
        self.tree.bind("<B1-Motion>", self.on_drag_motion)
        self.tree.bind("<ButtonRelease-1>", self.on_drag_drop)
        self.drag_data = {"item": None}

        scrollbar = ctk.CTkScrollbar(frame_tree, command=self.tree.yview)
        self.tree.configure(yscroll=scrollbar.set)
        
        self.tree.grid(row=0, column=0, sticky="nsew")
        scrollbar.grid(row=0, column=1, sticky="ns")

        # --- BLOCO INFERIOR (Ação) ---
        frame_bot = ctk.CTkFrame(parent, fg_color="transparent")
        frame_bot.grid(row=3, column=0, sticky="ew", pady=(10, 0))
        
        self.btn_rename = ctk.CTkButton(frame_bot, text="2. INICIAR PROCESSAMENTO PARALELO (PRO)", command=self.rename_files,
                                        fg_color=self.c_primary, hover_color="#1B5E20", text_color="white", font=ctk.CTkFont(size=14, weight="bold"), height=50, state="disabled")
        self.btn_rename.pack(fill="x")
        
        self.frame_progress = ctk.CTkFrame(frame_bot, fg_color="transparent")
        # Escondido por padrão
        
        self.lbl_status = ctk.CTkLabel(self.frame_progress, text="Aguardando...", font=ctk.CTkFont(weight="bold"))
        self.lbl_status.pack(anchor="w", pady=(10, 0))
        
        self.progress = ctk.CTkProgressBar(self.frame_progress, mode="determinate", progress_color=self.c_primary, height=20)
        self.progress.set(0)
        self.progress.pack(fill="x", pady=5)
        
        self.lbl_perc = ctk.CTkLabel(self.frame_progress, text="0%", font=ctk.CTkFont(weight="bold", size=14))
        self.lbl_perc.pack()

    def create_pdf_tab(self, parent):
        parent.columnconfigure(0, weight=1)
        parent.rowconfigure(1, weight=1)

        # --- Controles superiores ---
        frame_top = ctk.CTkFrame(parent, fg_color="transparent")
        frame_top.grid(row=0, column=0, sticky="ew", pady=(0, 5))
        frame_top.columnconfigure(1, weight=1)

        ctk.CTkLabel(frame_top, text="NOME DO ARQUIVO FINAL:", font=ctk.CTkFont(weight="bold")).grid(row=0, column=0, sticky="w", pady=8)
        ctk.CTkEntry(frame_top, textvariable=self.pdf_output_name, height=36).grid(row=0, column=1, sticky="ew", padx=10, pady=8)

        frame_btns = ctk.CTkFrame(frame_top, fg_color="transparent")
        frame_btns.grid(row=0, column=2, sticky="e")
        ctk.CTkButton(frame_btns, text="+ Adicionar PDFs", command=self.pdf_add_files, fg_color=self.c_secondary, hover_color="#311B92", height=36).pack(side="left", padx=(0,5))
        ctk.CTkButton(frame_btns, text="Limpar Lista", command=self.pdf_clear_list, fg_color="gray40", hover_color="gray25", height=36).pack(side="left")

        # --- TreeView de PDFs ---
        frame_tree = ctk.CTkFrame(parent, fg_color="transparent")
        frame_tree.grid(row=1, column=0, sticky="nsew", pady=8)
        frame_tree.columnconfigure(0, weight=1)
        frame_tree.rowconfigure(0, weight=1)

        style = ttk.Style()
        style.configure("PDF.Treeview",
                        rowheight=32, font=("Helvetica", 11),
                        background="#ffffff" if ctk.get_appearance_mode() == "Light" else "#2b2b2b",
                        foreground="#000000" if ctk.get_appearance_mode() == "Light" else "#ffffff",
                        fieldbackground="#ffffff" if ctk.get_appearance_mode() == "Light" else "#2b2b2b")
        style.configure("PDF.Treeview.Heading", font=("Helvetica", 11, "bold"),
                        background=self.c_primary, foreground="white")
        style.map("PDF.Treeview", background=[("selected", self.c_secondary)])

        pdf_cols = ("pos", "nome", "tamanho")
        self.pdf_tree = Treeview(frame_tree, columns=pdf_cols, show="headings", selectmode="browse", style="PDF.Treeview")
        self.pdf_tree.heading("pos", text="#")
        self.pdf_tree.heading("nome", text="ARQUIVO")
        self.pdf_tree.heading("tamanho", text="TAMANHO")
        self.pdf_tree.column("pos", width=50, anchor="center")
        self.pdf_tree.column("nome", anchor="w")
        self.pdf_tree.column("tamanho", width=100, anchor="center")

        self.pdf_tree.bind("<ButtonPress-1>", self.pdf_drag_start)
        self.pdf_tree.bind("<B1-Motion>", self.pdf_drag_motion)
        self.pdf_tree.bind("<ButtonRelease-1>", self.pdf_drag_drop)
        self.pdf_drag_data = {"item": None}

        pdf_scroll = ctk.CTkScrollbar(frame_tree, command=self.pdf_tree.yview)
        self.pdf_tree.configure(yscroll=pdf_scroll.set)
        self.pdf_tree.grid(row=0, column=0, sticky="nsew")
        pdf_scroll.grid(row=0, column=1, sticky="ns")

        # Botões de reordenar
        frame_order = ctk.CTkFrame(parent, fg_color="transparent")
        frame_order.grid(row=2, column=0, sticky="ew", pady=(0, 5))
        ctk.CTkButton(frame_order, text="⬆ Subir", command=self.pdf_move_up, fg_color="gray40", hover_color="gray25", width=100).pack(side="left", padx=(0,5))
        ctk.CTkButton(frame_order, text="⬇ Descer", command=self.pdf_move_down, fg_color="gray40", hover_color="gray25", width=100).pack(side="left")

        # --- Botão e Progresso ---
        frame_bot = ctk.CTkFrame(parent, fg_color="transparent")
        frame_bot.grid(row=3, column=0, sticky="ew", pady=(5, 0))

        self.pdf_merge_btn = ctk.CTkButton(frame_bot, text="UNIFICAR E COMPRIMIR PDFs", command=self.merge_pdfs,
                      fg_color=self.c_primary, text_color="white", font=ctk.CTkFont(size=15, weight="bold"), height=50, corner_radius=8)
        self.pdf_merge_btn.pack(fill="x")

        self.pdf_progress_frame = ctk.CTkFrame(frame_bot, fg_color="transparent")
        self.pdf_status_lbl = ctk.CTkLabel(self.pdf_progress_frame, text="Aguardando...", font=ctk.CTkFont(weight="bold"))
        self.pdf_status_lbl.pack(anchor="w", pady=(8,0))
        self.pdf_progress = ctk.CTkProgressBar(self.pdf_progress_frame, mode="determinate", progress_color=self.c_primary, height=18)
        self.pdf_progress.set(0)
        self.pdf_progress.pack(fill="x", pady=4)
        self.pdf_perc_lbl = ctk.CTkLabel(self.pdf_progress_frame, text="0%", font=ctk.CTkFont(weight="bold", size=13))
        self.pdf_perc_lbl.pack()

    # ---------- PDF TAB helpers ----------
    def pdf_add_files(self):
        files = filedialog.askopenfilenames(title="Selecione os PDFs", filetypes=[("PDF files", "*.pdf")])
        for f in files:
            if f not in [p['path'] for p in self.pdf_files]:
                size = os.path.getsize(f)
                self.pdf_files.append({'path': f, 'name': os.path.basename(f), 'size': size})
        self._refresh_pdf_tree()

    def pdf_clear_list(self):
        self.pdf_files.clear()
        self._refresh_pdf_tree()

    def _refresh_pdf_tree(self):
        for item in self.pdf_tree.get_children():
            self.pdf_tree.delete(item)
        for i, p in enumerate(self.pdf_files):
            size_str = self.format_size(p['size'])
            self.pdf_tree.insert("", "end", values=(i+1, p['name'], size_str))

    def pdf_drag_start(self, event):
        item = self.pdf_tree.identify_row(event.y)
        if item:
            self.pdf_drag_data["item"] = item
            self.pdf_tree.config(cursor="hand2")

    def pdf_drag_motion(self, event):
        pass

    def pdf_drag_drop(self, event):
        self.pdf_tree.config(cursor="")
        if not self.pdf_drag_data["item"]: return
        target = self.pdf_tree.identify_row(event.y)
        src = self.pdf_drag_data["item"]
        if target and target != src:
            si = self.pdf_tree.index(src)
            ti = self.pdf_tree.index(target)
            item = self.pdf_files.pop(si)
            self.pdf_files.insert(ti, item)
            self._refresh_pdf_tree()
        self.pdf_drag_data["item"] = None

    def pdf_move_up(self):
        sel = self.pdf_tree.selection()
        if not sel: return
        idx = self.pdf_tree.index(sel[0])
        if idx > 0:
            self.pdf_files[idx], self.pdf_files[idx-1] = self.pdf_files[idx-1], self.pdf_files[idx]
            self._refresh_pdf_tree()
            new_item = self.pdf_tree.get_children()[idx-1]
            self.pdf_tree.selection_set(new_item)

    def pdf_move_down(self):
        sel = self.pdf_tree.selection()
        if not sel: return
        idx = self.pdf_tree.index(sel[0])
        if idx < len(self.pdf_files) - 1:
            self.pdf_files[idx], self.pdf_files[idx+1] = self.pdf_files[idx+1], self.pdf_files[idx]
            self._refresh_pdf_tree()
            new_item = self.pdf_tree.get_children()[idx+1]
            self.pdf_tree.selection_set(new_item)

    def create_utils_tab(self, parent):
        # Frame scrollável para os cards
        parent.columnconfigure(0, weight=1)
        parent.columnconfigure(1, weight=1)

        ctk.CTkLabel(parent, text="CAIXA DE FERRAMENTAS AVANÇADAS",
                     font=ctk.CTkFont(size=20, weight="bold"), text_color=self.c_primary
                     ).grid(row=0, column=0, columnspan=2, pady=(15, 20))

        # Cada card tem: status_label próprio + botão que roda diálogo na main thread
        def build_card(col, row, icon, title, desc, btn_text, cmd):
            frame = ctk.CTkFrame(parent, corner_radius=15, fg_color="transparent",
                                 border_width=2, border_color="gray70")
            frame.grid(row=row, column=col, padx=12, pady=10, sticky="nsew")
            ctk.CTkLabel(frame, text=icon, font=ctk.CTkFont(size=32)).pack(pady=(12,0))
            ctk.CTkLabel(frame, text=title, font=ctk.CTkFont(size=15, weight="bold")).pack(pady=4)
            ctk.CTkLabel(frame, text=desc, font=ctk.CTkFont(size=11), text_color="gray",
                         wraplength=280, justify="center").pack(pady=(0,10), padx=10)
            status = ctk.CTkLabel(frame, text="⏳ Aguardando...",
                                  font=ctk.CTkFont(size=11), text_color="gray")
            status.pack(pady=(0,6))
            prog = ctk.CTkProgressBar(frame, mode="determinate",
                                      progress_color=self.c_primary, height=8)
            prog.set(0)
            prog.pack(fill="x", padx=16, pady=(0,8))
            btn = ctk.CTkButton(frame, text=btn_text, command=lambda c=cmd, s=status, p=prog: c(s, p),
                                fg_color=self.c_secondary, hover_color="#311B92",
                                font=ctk.CTkFont(weight="bold"))
            btn.pack(pady=(0,16), padx=16, fill="x")
            return frame

        build_card(0, 1, "🔤", "PDF para Word",
                   "Converte PDF com texto real para .docx totalmente editável.",
                   "Converter PDF → Word", self.do_pdf_to_word)
        build_card(1, 1, "📊", "Excel para PDF",
                   "Gera um relatório PDF bem formatado a partir de uma planilha .xlsx.",
                   "Converter Excel → PDF", self.do_excel_to_pdf)
        build_card(0, 2, "📑", "PDF para Excel",
                   "Extrai tabelas de PDFs nativos e exporta os dados em Excel.",
                   "Extrair Tabelas → Excel", self.do_pdf_to_excel)
        build_card(1, 2, "🖼️", "Imagens para PDF",
                   "Une múltiplas fotos JPG/PNG num único álbum PDF de alta qualidade.",
                   "Unir Fotos → PDF", self.do_jpg_to_pdf)
        build_card(0, 3, "✂️", "Dividir PDF",
                   "Separa cada página de um PDF em arquivos individuais numa pasta.",
                   "Separar Páginas", self.do_split_pdf)

    # ----------- Helpers thread-safe -----------
    def _show_info(self, title, msg):
        self.root.after(0, lambda: messagebox.showinfo(title, msg))

    def _show_err(self, title, msg):
        self.root.after(0, lambda: messagebox.showerror(title, msg))

    def _set_status(self, lbl, prog, text, pval=None):
        """Atualiza status e progresso de forma thread-safe."""
        def _do():
            lbl.configure(text=text)
            if pval is not None:
                prog.set(pval)
        self.root.after(0, _do)

    # ----------- Motor 1: PDF → Word -----------
    def do_pdf_to_word(self, status_lbl, prog):
        # Diálogo na thread PRINCIPAL
        file = filedialog.askopenfilename(title="Selecione o PDF de origem",
                                          filetypes=[("PDF", "*.pdf")])
        if not file: return
        out = filedialog.asksaveasfilename(title="Salvar Word como...",
                                           defaultextension=".docx",
                                           filetypes=[("Word", "*.docx")],
                                           initialfile=os.path.splitext(os.path.basename(file))[0] + ".docx")
        if not out: return
        threading.Thread(target=self._bg_pdf_to_word, args=(file, out, status_lbl, prog), daemon=True).start()

    def _bg_pdf_to_word(self, file, out, lbl, prog):
        try:
            self._set_status(lbl, prog, "🔄 Analisando PDF...", 0.1)
            # Verifica se o PDF tem texto extraível nativo
            texto_total = ""
            with pdfplumber.open(file) as pdf_check:
                for pg in pdf_check.pages:
                    texto_total += (pg.extract_text() or "")

            if len(texto_total.strip()) >= 30:
                # ── CAMINHO 1: PDF nativo com texto → pdf2docx (preserva layout)
                self._set_status(lbl, prog, "🔄 Convertendo PDF nativo para Word...", 0.4)
                cv = Converter(file)
                cv.convert(out)
                cv.close()
            else:
                # ── CAMINHO 2: PDF escaneado → OCR com Tesseract + python-docx
                if not _PYTESSERACT_OK:
                    self._set_status(lbl, prog, "❌ Tesseract não instalado!", 0.0)
                    self._show_err("Tesseract não encontrado",
                        "Para processar PDFs escaneados, o Tesseract OCR precisa estar instalado.\n\n"
                        "• Mac: abra o Terminal e execute:\n  brew install tesseract tesseract-lang\n\n"
                        "• Windows: baixe em github.com/UB-Mannheim/tesseract/wiki")
                    return
                self._set_status(lbl, prog, "🔍 PDF escaneado detectado! Iniciando OCR...", 0.1)
                tess_path = shutil.which("tesseract")
                if not tess_path:
                    for p in ["/opt/homebrew/bin/tesseract", "/usr/local/bin/tesseract",
                               r"C:\Program Files\Tesseract-OCR\tesseract.exe"]:
                        if os.path.exists(p):
                            tess_path = p
                            break
                if tess_path:
                    _pytesseract_module.pytesseract.tesseract_cmd = tess_path
                else:
                    self._set_status(lbl, prog, "❌ Tesseract não encontrado!", 0.0)
                    self._show_err("Tesseract não encontrado",
                        "• Mac: brew install tesseract tesseract-lang\n"
                        "• Windows: github.com/UB-Mannheim/tesseract/wiki")
                    return

                from docx import Document
                from docx.shared import Pt
                doc_word = Document()
                doc_word.add_heading("Documento EcoRenamer OCR", level=1)
                doc_fitz = fitz.open(file)
                total_pgs = len(doc_fitz)
                for i, page in enumerate(doc_fitz):
                    perc = 0.15 + (i / total_pgs) * 0.80
                    self._set_status(lbl, prog, f"🔍 OCR: Página {i+1}/{total_pgs}...", perc)
                    mat = fitz.Matrix(300/72, 300/72)
                    pix = page.get_pixmap(matrix=mat)
                    with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                        tmp_path = tmp.name
                    pix.save(tmp_path)
                    pil_img = Image.open(tmp_path)
                    try:
                        texto_ocr = _pytesseract_module.image_to_string(pil_img, lang="por+eng")
                    except Exception:
                        texto_ocr = _pytesseract_module.image_to_string(pil_img, lang="eng")
                    pil_img.close()
                    os.remove(tmp_path)
                    if i > 0:
                        doc_word.add_page_break()
                    p = doc_word.add_paragraph()
                    run = p.add_run(texto_ocr)
                    run.font.size = Pt(11)
                doc_fitz.close()
                doc_word.save(out)

            self._set_status(lbl, prog, "✅ Concluído!", 1.0)
            self._show_info("Sucesso", f"Documento Word editável salvo em:\n{out}")
        except Exception as e:
            self._set_status(lbl, prog, f"❌ Erro: {e}", 0.0)
            self._show_err("Erro", f"Falha na conversão PDF → Word:\n{e}")

    # ----------- Motor 2: Excel → PDF -----------
    def do_excel_to_pdf(self, status_lbl, prog):
        file = filedialog.askopenfilename(title="Selecione a planilha Excel",
                                          filetypes=[("Excel", "*.xlsx")])
        if not file: return
        out = filedialog.asksaveasfilename(title="Salvar PDF como...",
                                           defaultextension=".pdf",
                                           filetypes=[("PDF", "*.pdf")],
                                           initialfile=os.path.splitext(os.path.basename(file))[0] + ".pdf")
        if not out: return
        threading.Thread(target=self._bg_excel_to_pdf, args=(file, out, status_lbl, prog), daemon=True).start()

    def _bg_excel_to_pdf(self, file, out, lbl, prog):
        try:
            self._set_status(lbl, prog, "🔄 Lendo planilha...", 0.2)
            wb = openpyxl.load_workbook(file, data_only=True)
            sheet = wb.active
            data = []
            for row in sheet.iter_rows(values_only=True):
                cleaned = [str(c) if c is not None else "" for c in row]
                if any(c.strip() for c in cleaned):  # pula linhas totalmente vazias
                    data.append(cleaned)
            if not data:
                self._set_status(lbl, prog, "⚠️ Planilha vazia!", 0.0)
                self._show_err("Erro", "A planilha selecionada parece estar vazia.")
                return
            self._set_status(lbl, prog, "🔄 Gerando PDF...", 0.6)
            from reportlab.lib.units import cm
            col_count = max(len(r) for r in data)
            # Calcula largura de coluna disponível
            available_width = landscape(A4)[0] - 2*cm
            col_width = available_width / col_count
            col_widths = [col_width] * col_count
            pdf_doc = SimpleDocTemplate(out, pagesize=landscape(A4),
                                        leftMargin=cm, rightMargin=cm,
                                        topMargin=cm, bottomMargin=cm)
            table = Table(data, colWidths=col_widths, repeatRows=1)
            table_style = TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor(self.c_primary)),
                ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
                ('FONTSIZE', (0,0), (-1,0), 9),
                ('BOTTOMPADDING', (0,0), (-1,0), 8),
                ('TOPPADDING', (0,0), (-1,0), 8),
                ('BACKGROUND', (0,1), (-1,-1), colors.white),
                ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F5F5F5')]),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CCCCCC')),
                ('FONTNAME', (0,1), (-1,-1), 'Helvetica'),
                ('FONTSIZE', (0,1), (-1,-1), 8),
                ('WORDWRAP', (0,0), (-1,-1), True),
            ])
            table.setStyle(table_style)
            pdf_doc.build([table])
            self._set_status(lbl, prog, "✅ Concluído!", 1.0)
            self._show_info("Sucesso", f"Relatório PDF salvo em:\n{out}")
        except Exception as e:
            self._set_status(lbl, prog, f"❌ Erro: {e}", 0.0)
            self._show_err("Erro", f"Falha ao gerar PDF:\n{str(e)}")

    # ----------- Motor 3: PDF → Excel -----------
    def do_pdf_to_excel(self, status_lbl, prog):
        file = filedialog.askopenfilename(title="Selecione o PDF com tabelas",
                                          filetypes=[("PDF", "*.pdf")])
        if not file: return
        out = filedialog.asksaveasfilename(title="Salvar Excel como...",
                                           defaultextension=".xlsx",
                                           filetypes=[("Excel", "*.xlsx")],
                                           initialfile=os.path.splitext(os.path.basename(file))[0] + "_Extraido.xlsx")
        if not out: return
        threading.Thread(target=self._bg_pdf_to_excel, args=(file, out, status_lbl, prog), daemon=True).start()

    def _bg_pdf_to_excel(self, file, out, lbl, prog):
        try:
            self._set_status(lbl, prog, "🔄 Analisando PDF...", 0.1)
            wb = openpyxl.Workbook()
            ws = wb.active
            tabelas_encontradas = 0
            with pdfplumber.open(file) as pdf:
                total_pgs = len(pdf.pages)
                for i, page in enumerate(pdf.pages):
                    self._set_status(lbl, prog, f"🔄 Página {i+1}/{total_pgs}...", (i+1)/total_pgs * 0.85)
                    # Tenta extrair tabela estruturada
                    table = page.extract_table()
                    if table:
                        tabelas_encontradas += 1
                        if tabelas_encontradas > 1:
                            ws.append([])  # linha em branco entre tabelas
                            ws.append([f"--- Página {i+1} ---"])
                        for row in table:
                            ws.append([str(c).strip() if c is not None else "" for c in row])
                    else:
                        # Fallback: extrai texto bruto linha a linha
                        texto = page.extract_text()
                        if texto:
                            tabelas_encontradas += 1
                            ws.append([f"--- Página {i+1} (texto) ---"])
                            for linha in texto.split("\n"):
                                ws.append([linha])
            if tabelas_encontradas == 0:
                self._set_status(lbl, prog, "⚠️ Nenhuma tabela encontrada!", 0.0)
                self._show_err("Aviso", "Nenhum conteúdo extraível foi encontrado neste PDF.\nVerifique se ele possui texto ou tabelas nativas (não escaneado).")
                return
            wb.save(out)
            self._set_status(lbl, prog, "✅ Concluído!", 1.0)
            self._show_info("Sucesso", f"{tabelas_encontradas} página(s) extraída(s) para:\n{out}")
        except Exception as e:
            self._set_status(lbl, prog, f"❌ Erro: {e}", 0.0)
            self._show_err("Erro", f"Falha na extração:\n{str(e)}")

    # ----------- Motor 4: Fotos → PDF -----------
    def do_jpg_to_pdf(self, status_lbl, prog):
        files = filedialog.askopenfilenames(
            title="Selecione as imagens (a ordem de seleção = ordem no PDF)",
            filetypes=[("Imagens", "*.jpg *.jpeg *.png *.bmp *.tiff")])
        if not files: return
        out = filedialog.asksaveasfilename(title="Salvar PDF como...",
                                           defaultextension=".pdf",
                                           filetypes=[("PDF", "*.pdf")],
                                           initialfile="Album_Imagens.pdf")
        if not out: return
        threading.Thread(target=self._bg_jpg_to_pdf, args=(list(files), out, status_lbl, prog), daemon=True).start()

    def _bg_jpg_to_pdf(self, files, out, lbl, prog):
        try:
            self._set_status(lbl, prog, "🔄 Processando imagens...", 0.05)
            doc = fitz.open()
            total = len(files)
            for i, f in enumerate(files):
                self._set_status(lbl, prog, f"🔄 Imagem {i+1}/{total}", (i+1)/total * 0.9)
                # Corrige rotação EXIF antes de inserir
                pil_img = Image.open(f)
                pil_img = ImageOps.exif_transpose(pil_img)
                pil_img = pil_img.convert("RGB")
                # Salva temporariamente como JPEG sem metadados conflitantes
                import tempfile
                with tempfile.NamedTemporaryFile(suffix=".jpg", delete=False) as tmp:
                    tmp_path = tmp.name
                pil_img.save(tmp_path, "JPEG", quality=90)
                pil_img.close()
                # Insere no PDF via PyMuPDF (muito mais robusto)
                img_doc = fitz.open(tmp_path)
                pdfbytes = img_doc.convert_to_pdf()
                img_doc.close()
                os.remove(tmp_path)
                img_pdf = fitz.open("pdf", pdfbytes)
                doc.insert_pdf(img_pdf)
                img_pdf.close()
            doc.save(out, deflate=True)
            doc.close()
            self._set_status(lbl, prog, "✅ Concluído!", 1.0)
            self._show_info("Sucesso", f"{total} imagem(ns) unidas em:\n{out}")
        except Exception as e:
            self._set_status(lbl, prog, f"❌ Erro: {e}", 0.0)
            self._show_err("Erro", f"Falha ao gerar PDF de imagens:\n{str(e)}")

    # ----------- Motor 5: Dividir PDF -----------
    def do_split_pdf(self, status_lbl, prog):
        file = filedialog.askopenfilename(title="Selecione o PDF para dividir",
                                          filetypes=[("PDF", "*.pdf")])
        if not file: return
        dest_folder = filedialog.askdirectory(title="Escolha a pasta de destino das páginas")
        if not dest_folder: return
        threading.Thread(target=self._bg_split_pdf, args=(file, dest_folder, status_lbl, prog), daemon=True).start()

    def _bg_split_pdf(self, file, dest_folder, lbl, prog):
        try:
            base_nome = os.path.splitext(os.path.basename(file))[0]
            doc = fitz.open(file)
            total = len(doc)
            self._set_status(lbl, prog, f"🔄 Dividindo {total} páginas...", 0.05)
            for i in range(total):
                self._set_status(lbl, prog, f"🔄 Página {i+1}/{total}", (i+1)/total * 0.95)
                novo_doc = fitz.open()
                novo_doc.insert_pdf(doc, from_page=i, to_page=i)
                novo_out = os.path.join(dest_folder, f"{base_nome}_Pag_{i+1:03d}.pdf")
                novo_doc.save(novo_out, deflate=True)
                novo_doc.close()
            doc.close()
            self._set_status(lbl, prog, "✅ Concluído!", 1.0)
            self._show_info("Sucesso", f"{total} páginas salvas em:\n{dest_folder}")
        except Exception as e:
            self._set_status(lbl, prog, f"❌ Erro: {e}", 0.0)
            self._show_err("Erro", f"Erro ao dividir PDF:\n{str(e)}")

    # ------------------ DRAG & DROP ------------------
    def on_drag_start(self, event):
        item = self.tree.identify_row(event.y)
        if item:
            self.drag_data["item"] = item
            self.tree.config(cursor="hand2")

    def on_drag_motion(self, event):
        pass

    def on_drag_drop(self, event):
        self.tree.config(cursor="")
        if not self.drag_data["item"]: return
            
        target_item = self.tree.identify_row(event.y)
        source_item = self.drag_data["item"]
        
        if target_item and target_item != source_item:
            source_idx = self.tree.index(source_item)
            target_idx = self.tree.index(target_item)
            
            moved_item = self.mapping.pop(source_idx)
            self.mapping.insert(target_idx, moved_item)
            
            self.update_mapping_after_reorder()
            
        self.drag_data["item"] = None

    def update_mapping_after_reorder(self):
        if not hasattr(self, 'original_dest_bases') or not self.original_dest_bases:
            return

        try:
            digitos = int(self.digits.get())
        except:
            digitos = 2
            
        folder = self.img_folder.get()
        compress = self.compress_var.get()

        for i, item in enumerate(self.mapping):
            if i < len(self.original_dest_bases):
                img_val_str = self.original_dest_bases[i]
                
                if img_val_str.isdigit(): novo_base = img_val_str.zfill(digitos)
                else: novo_base = img_val_str
                
                _, ext = os.path.splitext(item['orig_name'])
                if compress: item['new_name'] = f"{novo_base}.JPG"
                else: item['new_name'] = f"{novo_base}{ext.upper()}"
                
                item['new_path'] = os.path.join(folder, item['new_name'])

        for item_id in self.tree.get_children():
            self.tree.delete(item_id)
            
        for idx, item in enumerate(self.mapping):
            self.tree.insert("", "end", values=(idx + 1, item['orig_name'], item['new_name'], item['size_orig_str'], item['size_est_str']))

    # ------------------ LÓGICA IMAGENS ------------------
    def format_size(self, size_bytes):
        if size_bytes >= 1024 * 1024:
            return f"{size_bytes / (1024 * 1024):.2f} MB"
        else:
            return f"{size_bytes / 1024:.0f} KB"

    def select_img_folder(self):
        folder = filedialog.askdirectory(title="Selecione a pasta com as fotos")
        if folder:
            self.img_folder.set(folder)
            self.reset_preview()

    def select_excel(self):
        file = filedialog.askopenfilename(title="Selecione o arquivo Excel", filetypes=[("Excel files", "*.xlsx")])
        if file:
            self.excel_file.set(file)
            self.reset_preview()

    def reset_preview(self):
        self.mapping = []
        for item in self.tree.get_children():
            self.tree.delete(item)
        self.btn_rename.configure(state="disabled")

    def load_data(self):
        folder = self.img_folder.get()
        excel_path = self.excel_file.get()
        
        if not folder:
            messagebox.showwarning("Aviso", "Por favor, selecione a pasta de imagens.")
            return

        valid_exts = {".jpg", ".jpeg", ".png", ".heic"}
        try:
            arquivos = os.listdir(folder)
            images = [f for f in arquivos if not f.startswith('.') and os.path.splitext(f)[1].lower() in valid_exts]
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao acessar pasta:\n{e}")
            return
            
        if not images:
            messagebox.showwarning("Aviso", "Nenhuma imagem válida encontrada.")
            return
            
        ordem = self.sort_order.get()
        if "Decrescente" in ordem: images.sort(reverse=True)
        else: images.sort(reverse=False)
        
        excel_img_values = []
        if excel_path:
            try:
                wb = openpyxl.load_workbook(excel_path, data_only=True)
                sheet = wb.active
                img_col_idx = None
                for col_idx in range(1, sheet.max_column + 1):
                    cell_value = sheet.cell(row=1, column=col_idx).value
                    if cell_value and str(cell_value).strip().upper() == "IMG":
                        img_col_idx = col_idx; break
                if img_col_idx is None:
                    messagebox.showerror("Erro", "Coluna 'IMG' não encontrada.")
                    return
                for row_idx in range(2, sheet.max_row + 1):
                    if row_idx in sheet.row_dimensions and sheet.row_dimensions[row_idx].hidden: continue
                    val = sheet.cell(row=row_idx, column=img_col_idx).value
                    if val is not None: excel_img_values.append(val)
            except Exception as e:
                messagebox.showerror("Erro", f"Erro no Excel:\n{e}")
                return
            if not excel_img_values:
                messagebox.showwarning("Aviso", "Coluna 'IMG' está vazia.")
                return
            if len(images) != len(excel_img_values):
                messagebox.showwarning("Aviso", f"Imagens ({len(images)}) vs Excel ({len(excel_img_values)}).")
            qtd_pares = min(len(images), len(excel_img_values))
        else:
            qtd_pares = len(images)
            try: start_no = int(self.start_number.get())
            except: start_no = 1
            excel_img_values = [str(start_no + i) for i in range(qtd_pares)]
            
        self.mapping = []
        self.original_dest_bases = []
        try: digitos = int(self.digits.get())
        except: digitos = 2
        
        seen_targets = set()
        duplicados = set()
        
        for i in range(qtd_pares):
            orig_name = images[i]
            img_val_str = str(excel_img_values[i]).strip()
            if img_val_str.endswith(".0"): img_val_str = img_val_str[:-2]
            if img_val_str.isdigit(): novo_base = img_val_str.zfill(digitos)
            else: novo_base = img_val_str
            self.original_dest_bases.append(img_val_str)
                
            _, ext = os.path.splitext(orig_name)
            if self.compress_var.get(): novo_nome = f"{novo_base}.JPG"
            else: novo_nome = f"{novo_base}{ext.upper()}"
            
            if novo_nome in seen_targets: duplicados.add(novo_nome)
            seen_targets.add(novo_nome)
            
            orig_path = os.path.join(folder, orig_name)
            novo_path = os.path.join(folder, novo_nome)
            
            self.mapping.append({
                'orig_name': orig_name, 'new_name': novo_nome,
                'orig_path': orig_path, 'new_path': novo_path
            })
            
        if duplicados:
            messagebox.showwarning("Atenção", f"Conflitos de nome previstos (Ex: {list(duplicados)[0]}). O novo motor lida com swaps, mas cuidado com cópias puras.")
            
        for item in self.tree.get_children(): self.tree.delete(item)
        for idx, item in enumerate(self.mapping):
            item['size_orig_str'] = self.format_size(os.path.getsize(item['orig_path']))
            item['size_est_str'] = "70 KB" 
            self.tree.insert("", "end", values=(idx + 1, item['orig_name'], item['new_name'], item['size_orig_str'], item['size_est_str']))
            
        if self.mapping:
            self.btn_rename.configure(state="normal")
            
    def rename_files(self):
        if not self.mapping or self.processing: return
        if not messagebox.askyesno("Confirmar", f"Renomear {len(self.mapping)} arquivos?"): return
            
        self.processing = True
        self.btn_rename.configure(state="disabled")
        self.frame_progress.pack(fill="x", pady=10)
        self.progress.set(0)
        self.lbl_status.configure(text="Iniciando motor turbo (ThreadPool)...")
        
        threading.Thread(target=self.run_rename_task_robust, daemon=True).start()

    def run_rename_task_robust(self):
        sucessos = 0
        falhas = 0
        erros_msg = []
        mapping_sq = list(self.mapping)
        total = len(mapping_sq)
        
        from concurrent.futures import ThreadPoolExecutor, as_completed
        
        with ThreadPoolExecutor() as executor:
            # Submete todas as tarefas
            futures = {executor.submit(self.process_single_image, item, idx): item for idx, item in enumerate(mapping_sq)}
            
            concluidos = 0
            for future in as_completed(futures):
                concluidos += 1
                res_ok, res_msg = future.result()
                if res_ok: sucessos += 1
                else: 
                    falhas += 1
                    erros_msg.append(res_msg)
                
                perc_num = concluidos / total
                # ATUALIZAÇÃO SEGURA DA GUI: Passar para a thread principal (evita hang/deadlock do Tkinter)
                self.root.after(0, self.update_ui_progress, perc_num, int(perc_num * 100), f"Fase 1/2: Compactando ({concluidos}/{total})")

        self.root.after(0, self.lbl_status.configure, {"text": "Fase 2/2: Aplicando nomes definitivos..."})
        for item in mapping_sq:
            temp_path = item['new_path'] + ".ecotmp"
            if os.path.exists(temp_path):
                try:
                    if os.path.exists(item['orig_path']): os.remove(item['orig_path'])
                    if os.path.exists(item['new_path']): os.remove(item['new_path'])
                    os.rename(temp_path, item['new_path'])
                except Exception as e:
                    print(f"Erro cleanup {item['new_name']}: {e}")

        self.root.after(0, lambda: self.finish_rename(sucessos, falhas, erros_msg))

    def update_ui_progress(self, val_float, perc_int, status):
        # Esta função agora é chamada de forma segura pela thread principal via root.after
        self.progress.set(val_float)
        self.lbl_perc.configure(text=f"{perc_int}%")
        self.lbl_status.configure(text=status)

    def process_single_image(self, item, index):
        temp_target = item['new_path'] + ".ecotmp"
        try:
            if self.compress_var.get():
                img = Image.open(item['orig_path'])
                img = ImageOps.exif_transpose(img)
                img.thumbnail((800, 800), Image.Resampling.LANCZOS)
                img = img.convert("RGB")
                img.save(temp_target, "JPEG", optimize=True, quality=45)
                img.close()
            else:
                shutil.copy2(item['orig_path'], temp_target)
            return True, ""
        except OSError as e:
            err_str = str(e).lower()
            if hasattr(e, 'errno') and e.errno == 89 or "operation canceled" in err_str or "canceled" in err_str:
                return False, f"{item['orig_name']}: Preso na Nuvem! Espere o Google Drive/iCloud baixar a foto para o Mac."
            return False, f"{item['orig_name']}: Erro de disco: {e}"
        except Exception as e:
            return False, f"{item['orig_name']}: {str(e)}"

    def finish_rename(self, sucessos, falhas, erros_msg):
        msg = f"Processo concluído!\nSucesso: {sucessos}\nFalhas: {falhas}"
        if falhas > 0:
            messagebox.showwarning("Aviso", msg + f"\n\nErros:\n" + "\n".join(erros_msg[:5]))
        else:
            messagebox.showinfo("Sucesso Total", msg)
        
        self.processing = False
        self.frame_progress.pack_forget()
        self.reset_preview()

    # ------------------ PDFs ------------------
    def merge_pdfs(self):
        if not self.pdf_files:
            messagebox.showwarning("Aviso", "Adicione pelo menos um PDF à lista antes de unificar.")
            return

        out_name = self.pdf_output_name.get().strip() or "Documento_Unificado.pdf"
        if not out_name.endswith(".pdf"): out_name += ".pdf"

        out_path = filedialog.asksaveasfilename(
            title="Salvar PDF unificado como...",
            defaultextension=".pdf",
            filetypes=[("PDF", "*.pdf")],
            initialfile=out_name)
        if not out_path: return

        self.pdf_merge_btn.configure(state="disabled")
        self.pdf_progress_frame.pack(fill="x", pady=8)
        self.pdf_progress.set(0)
        self.pdf_status_lbl.configure(text="Iniciando unificação...")

        threading.Thread(target=self._bg_merge_pdfs, args=(out_path,), daemon=True).start()

    def _bg_merge_pdfs(self, out_path):
        total = len(self.pdf_files)
        try:
            doc_final = fitz.open()
            for i, p in enumerate(self.pdf_files):
                perc = (i + 1) / total
                self.root.after(0, lambda v=perc, i=i: (
                    self.pdf_progress.set(v),
                    self.pdf_perc_lbl.configure(text=f"{int(v*100)}%"),
                    self.pdf_status_lbl.configure(text=f"Processando: {self.pdf_files[i]['name']}")
                ))
                try:
                    doc_temp = fitz.open(p['path'])
                    doc_final.insert_pdf(doc_temp)
                    doc_temp.close()
                except Exception as e:
                    print(f"Erro ao ler {p['name']}: {e}")

            self.root.after(0, lambda: self.pdf_status_lbl.configure(text="💾 Salvando arquivo final..."))
            doc_final.save(out_path, garbage=4, deflate=True)
            doc_final.close()

            self.root.after(0, lambda: (
                self.pdf_progress.set(1.0),
                self.pdf_perc_lbl.configure(text="100%"),
                self.pdf_status_lbl.configure(text="✅ Concluído!"),
                self.pdf_merge_btn.configure(state="normal"),
                messagebox.showinfo("Sucesso", f"PDFs unificados com sucesso em:\n{out_path}")
            ))
        except Exception as e:
            self.root.after(0, lambda: (
                self.pdf_status_lbl.configure(text=f"❌ Erro: {e}"),
                self.pdf_merge_btn.configure(state="normal"),
                messagebox.showerror("Erro", str(e))
            ))

    def check_for_updates(self):
        try:
            with urllib.request.urlopen(UPDATE_URL, timeout=5) as r:
                data = json.loads(r.read().decode())
                remote_ver = data.get("version", VERSION)
                if remote_ver > VERSION:
                    msg = f"Nova versão encontrada: v{remote_ver}\nMudanças: {data.get('changelog')}\n\nDeseja fechar, instalar silenciosamente e reiniciar agora?"
                    if messagebox.askyesno("Atualização Pronta", msg):
                        os_name = platform.system()
                        url = data.get("download_url_mac") if os_name == "Darwin" else data.get("download_url_win")
                        if url: threading.Thread(target=self.run_auto_update, args=(url, remote_ver), daemon=True).start()
                        else: messagebox.showinfo("Erro", "URL de instalação não encontrada no servidor.")
                else: messagebox.showinfo("App em Dia", "Você já está rodando a última Enterprise Edition.")
        except: pass

    def run_auto_update(self, url, version):
        # UI Freeze de Download
        self.btn_rename.configure(state="disabled")
        try: self.btn_load.configure(state="disabled")
        except: pass
        self.frame_progress.pack(fill="x", pady=10)
        self.lbl_status.configure(text=f"Abaixando pacote v{version} (Em Background)...")
        self.progress.set(0.3)
        self.lbl_perc.configure(text="30%")
        
        try:
            temp_dir = os.path.join(os.path.expanduser("~"), ".ecowave_update_tmp")
            if os.path.exists(temp_dir): shutil.rmtree(temp_dir)
            os.makedirs(temp_dir)
            
            zpath = os.path.join(temp_dir, "u.zip")
            urllib.request.urlretrieve(url, zpath)
            
            self.lbl_status.configure(text="Pré-Instalando e Descompactando...")
            self.progress.set(0.8)
            with zipfile.ZipFile(zpath, 'r') as zf: zf.extractall(temp_dir)
            
            # Bloqueio caso seja rodado via Python Raw
            is_frozen = getattr(sys, 'frozen', False)
            if not is_frozen:
                messagebox.showinfo("Dev Mode", "Instalador baixado, mas ignorado pois não é um executável compilado (.app ou .exe).")
                self.frame_progress.pack_forget()
                return
                
            os_name = platform.system()
            if os_name == "Darwin":
                # APLICANDO TRAVA TÉCNICA E CORREÇÃO DE CAMINHO (A VIDA DOS SEUS ARQUIVOS)
                # O executável do Mac fica em: MeuApp.app/Contents/MacOS/app
                # Recuar 2 pastas (..) cai exatamente na pasta .app
                app_path = os.path.abspath(os.path.join(os.path.dirname(sys.executable), "..", ".."))
                new_app_extracted = os.path.join(temp_dir, "RenomeadorApp.app")
                
                # TRAVA DE SEGURANÇA: NUNCA subscrever se não for um .app real
                if not app_path.endswith(".app"):
                    messagebox.showerror("Erro de Segurança Crítico", f"O atualizador tentou sobrescrever um diretório inválido:\n{app_path}\nAtualização Abortada pelo Sistema de Defesa.")
                    self.frame_progress.pack_forget()
                    return
                
                script_sh = os.path.join(temp_dir, "updater.command")
                with open(script_sh, "w") as f:
                    f.write("#!/bin/bash\n")
                    f.write("sleep 2\n") # Tempo pro App fechar e liberar a pasta
                    # FLAG --delete FOI REMOVIDA PARA SEMPRE! Apenas substitui novos arquivos sem excluir nada em volta.
                    f.write(f"rsync -a '{new_app_extracted}/' '{app_path}/'\n") 
                    f.write(f"xattr -cr '{app_path}'\n") # BURACO NO GATEKEEPER: Remove quarentena
                    f.write(f"open '{app_path}'\n") # Reinicia
                    f.write(f"rm -rf '{temp_dir}'\n") # Limpeza ninja do zips
                    f.write("rm -- \"$0\"\n") # Auto-destrói o script sh
                    
                os.chmod(script_sh, 0o755)
                subprocess.Popen(["/bin/bash", script_sh], start_new_session=True)
                sys.exit(0) # Morre
                
            else: # Windows
                app_path = os.path.dirname(sys.executable)
                exe_name = os.path.basename(sys.executable)
                
                script_bat = os.path.join(temp_dir, "updater.bat")
                with open(script_bat, "w") as f:
                    f.write("@echo off\n")
                    f.write("timeout /t 2 /nobreak >nul\n") # Espera App fechar
                    f.write(f"xcopy /S /E /Y /I \"{temp_dir}\\*\" \"{app_path}\\\"\n") # Subscreve forçado
                    f.write(f"start \"\" \"{os.path.join(app_path, exe_name)}\"\n") # Inicia novo executável
                    f.write("del \"%~f0\"\n") # Pede pro arquivo bat tentar se matar
                    
                # Roda o BAT desanexado do nosso processo
                subprocess.Popen([script_bat], creationflags=subprocess.CREATE_NEW_CONSOLE)
                sys.exit(0) # Mata o App Atual
                
        except Exception as e:
            messagebox.showerror("Erro Crítico de Instalação", f"Falha ao realizar 'Seamless Update':\n\n{e}")
            self.frame_progress.pack_forget()


if __name__ == "__main__":
    # CustomTkinter precisa ser instanciado como ctk.CTk()
    root = ctk.CTk()
    app = ToolApp(root)
    root.mainloop()
